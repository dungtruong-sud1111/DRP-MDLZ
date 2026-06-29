#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Convert markdown TO-BE Blueprint -> .docx theo format mặc định (chỉnh theo brand khách).

Dùng:
    python md_to_docx.py "_TOBE-BLUEPRINT.md" [-o output.docx]

Format áp dụng: Times New Roman 12; header bảng nền #176bb4 (mặc định, đổi qua HEADER_FILL) chữ trắng đậm;
lề 3/2/3/2.5cm; footer số trang phải; heading map theo cấp #.
Mermaid: render PNG nếu có 'mmdc' (mermaid-cli), không thì để placeholder.
Tự bỏ phần "MỤC LỤC KẾ HOẠCH" (tracking).

TRANG BÌA (tùy chọn): đặt khối <!-- COVER ... --> ở ĐẦU file .md, script tự
dựng trang bìa (hàng 2 logo căn giữa → tên công ty → tên tài liệu → ngày ghim
ở CHÂN trang bìa) rồi ngắt sang trang nội dung. Cú pháp:
    <!-- COVER
    company: Công ty Cổ phần ...
    title: TÀI LIỆU GIẢI PHÁP ... (TO-BE) | TÊN MODULE ...
    date: 23/06/2026
    logo-left: logo-smartlog.png      # đường dẫn tương đối so với file .md
    logo-right: logo-khach.png        # logo thiếu -> ô placeholder
    drop-title: true                  # bỏ tiêu đề/metadata lặp ở đầu nội dung
    -->
Dùng | trong 'title' để xuống dòng. Logo .webp tự convert sang .png nếu có Pillow.
Ngày được ghim ở footer-trang-đầu nên LUÔN nằm đáy trang bìa (không tràn trang).

Yêu cầu: pip install python-docx   (mermaid-cli, Pillow là tùy chọn)
Luôn chạy với PYTHONIOENCODING=utf-8 PYTHONUTF8=1.
"""
import sys, os, re, shutil, subprocess, tempfile

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("[LỖI] Thiếu python-docx. Cài: pip install python-docx")
    sys.exit(1)

HEADER_FILL = "176BB4"  # màu header bảng mặc định — đổi theo brand khách
COVER_NAVY = RGBColor(0x1F, 0x38, 0x64)
COVER_GREY = RGBColor(0x80, 0x80, 0x80)
COVER_RE = re.compile(r"<!--\s*COVER\b(.*?)-->", re.DOTALL)
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
MH_RE = re.compile(r"\[\[\s*(MH-[^\]]+?)\s*\]\]")
IMG_RE = re.compile(r"!\[[^\]]*\]\(([^)]+)\)")

# Truy vết nội bộ — bỏ khi convert sang Word (giao khách)
# Khớp: _(Truy vết ...)_, _(Ghi chú ...)_, (Truy vết ...) ở cuối câu/đoạn
_TRACEBACK_RE = re.compile(
    r"_\s*\((?:Truy vết|Ghi chú|Nguồn)[^)]*\)\s*_"   # _(...truy vết...)_
    r"|(?<!\w)\((?:Truy vết|truy vết)\s[^)]+\)[,.]?\s*$",  # (Truy vết OBJ-xx) cuối câu
    re.MULTILINE,
)

def strip_traceback(text: str) -> str:
    """Bỏ annotation truy vết nội bộ khỏi text trước khi ghi Word."""
    return _TRACEBACK_RE.sub("", text).strip()


# ---------- helpers ----------
def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_inline(paragraph, text, base_bold=False, base_italic=False, color=None):
    """Thêm text vào paragraph, hiểu **đậm** và *nghiêng* cơ bản."""
    # tách theo bold trước
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            _add_italic_segment(paragraph, text[pos:m.start()], base_bold, base_italic, color)
        _add_italic_segment(paragraph, m.group(1), True, base_italic, color)
        pos = m.end()
    if pos < len(text):
        _add_italic_segment(paragraph, text[pos:], base_bold, base_italic, color)


def _add_italic_segment(paragraph, text, bold, italic, color):
    pos = 0
    for m in ITALIC_RE.finditer(text):
        if m.start() > pos:
            _run(paragraph, text[pos:m.start()], bold, italic, color)
        _run(paragraph, m.group(1), bold, True, color)
        pos = m.end()
    if pos < len(text):
        _run(paragraph, text[pos:], bold, italic, color)


def _run(paragraph, text, bold, italic, color):
    if text == "":
        return
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    if color is not None:
        r.font.color.rgb = color


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    fldStart = OxmlElement("w:fldChar"); fldStart.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldStart); run._r.append(instr); run._r.append(fldEnd)


def setup_document(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    # font cho chữ Đông Á (đảm bảo TNR áp cho mọi ký tự)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2.5)
        add_page_number_footer(section)


def add_heading(doc, text, level):
    try:
        p = doc.add_heading("", level=min(level, 9))
    except Exception:
        p = doc.add_paragraph()
    add_inline(p, text, base_bold=True)
    # Mỗi mục lớn (heading cấp 1 "# N.") bắt đầu ở đầu trang mới
    if level == 1:
        p.paragraph_format.page_break_before = True
    return p


def _find_swimlane_png(mmd_code: str, md_dir: str) -> str | None:
    """Tìm PNG swimlane tương ứng với block mermaid.
    Logic: đọc tên file .mmd đã lưu trong luu-do/ (dòng đầu comment # src: <name>)
    hoặc dò theo nội dung flowchart. Trả về đường dẫn PNG nếu tồn tại."""
    # Ưu tiên: dòng comment đầu tiên dạng "# src: in-in-bkd" hoặc "%% src: in-in-bkd"
    for line in mmd_code.splitlines():
        m = re.match(r"(?:#|%%)\s*src:\s*(.+)", line.strip())
        if m:
            name = m.group(1).strip()
            png = os.path.join(md_dir, "luu-do", name + ".png")
            if os.path.exists(png):
                return png

    # Fallback: quét tất cả PNG trong luu-do/ — dùng file JSON tương ứng để sinh nếu cần
    luu_do = os.path.join(md_dir, "luu-do")
    if not os.path.isdir(luu_do):
        return None

    # Tìm JSON có title khớp (heuristic: dùng file JSON đầu tiên tìm thấy nếu chỉ có 1)
    jsons = [f for f in os.listdir(luu_do) if f.endswith(".json")]
    if len(jsons) == 1:
        name = os.path.splitext(jsons[0])[0]
        png = os.path.join(luu_do, name + ".png")
        json_path = os.path.join(luu_do, jsons[0])
        if os.path.exists(png):
            return png
        # Chưa có PNG → sinh mới
        _generate_swimlane(json_path)
        if os.path.exists(png):
            return png
    return None


def _generate_swimlane(json_path: str):
    """Gọi swimlane_generator để sinh HTML + PNG từ JSON."""
    script = os.path.join(os.path.dirname(__file__), "swimlane_generator.py")
    if not os.path.exists(script):
        print(f"[CẢNH BÁO] Không tìm thấy swimlane_generator.py tại {script}")
        return
    try:
        subprocess.run(
            [sys.executable, script, json_path],
            check=True, timeout=120,
            env={**os.environ, "PYTHONIOENCODING": "utf-8", "PYTHONUTF8": "1"}
        )
    except Exception as e:
        print(f"[CẢNH BÁO] Sinh swimlane lỗi: {e}")


def render_mermaid(code, doc, mmdc, md_dir=None):
    """Ưu tiên dùng PNG swimlane nếu có; fallback mermaid-cli; fallback placeholder."""
    from docx.shared import Cm

    # 1. Tìm PNG swimlane tương ứng
    if md_dir:
        png_path = _find_swimlane_png(code, md_dir)
        if png_path and os.path.exists(png_path):
            try:
                # Chèn ảnh vừa khít chiều rộng trang (15.5cm = A4 trừ lề)
                doc.add_picture(png_path, width=Cm(15.5))
                print(f"[OK] Chèn swimlane PNG: {png_path}")
                return
            except Exception as e:
                print(f"[CẢNH BÁO] Chèn PNG lỗi ({e}); thử mermaid-cli.")

    # 2. Fallback: mermaid-cli
    if mmdc:
        try:
            tmpdir = tempfile.mkdtemp()
            mmd = os.path.join(tmpdir, "d.mmd")
            png = os.path.join(tmpdir, "d.png")
            with open(mmd, "w", encoding="utf-8") as f:
                f.write(code)
            subprocess.run([mmdc, "-i", mmd, "-o", png, "-b", "white"],
                           check=True, capture_output=True, timeout=120)
            if os.path.exists(png):
                doc.add_picture(png)
                return
        except Exception as e:
            print(f"[CẢNH BÁO] Render mermaid lỗi ({e}); để placeholder.")

    # 3. Placeholder
    p = doc.add_paragraph()
    _run(p, "[Lưu đồ — chèn ảnh thủ công hoặc dùng file .mmd]",
         bold=False, italic=True, color=RGBColor(0x80, 0x80, 0x80))


def add_table(doc, rows, banner=None):
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncol)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    # Dòng banner tiêu đề (gộp toàn bộ cột, nền navy đậm, chữ trắng) — khi heading đứng ngay trên bảng
    if banner:
        brow = table.add_row()
        merged = brow.cells[0]
        for ci in range(1, ncol):
            merged = merged.merge(brow.cells[ci])
        set_cell_shading(merged, HEADER_FILL)
        mp = merged.paragraphs[0]
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(mp, banner, base_bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for ri, cells in enumerate(rows):
        row = table.add_row()
        for ci in range(ncol):
            cell = row.cells[ci]
            text = cells[ci] if ci < len(cells) else ""
            cell.text = ""
            para = cell.paragraphs[0]
            is_colheader = (ri == 0)
            # có banner: dòng cột đậm chữ đen, nền trắng; không banner: dòng cột nền navy chữ trắng (như cũ)
            white = RGBColor(0xFF, 0xFF, 0xFF) if (is_colheader and not banner) else None
            segs = text.split("<br>")
            for si, seg in enumerate(segs):
                if si > 0:
                    para = cell.add_paragraph()
                add_inline(para, seg.strip(), base_bold=is_colheader, color=white)
            if is_colheader and not banner:
                set_cell_shading(cell, HEADER_FILL)


def is_table_sep(line):
    s = line.strip().strip("|")
    return bool(s) and all(re.match(r"^\s*:?-+:?\s*$", c) for c in s.split("|"))


def parse_table_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


# ---------- trang bìa ----------
def _set_cell_dashed_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "dashed"); e.set(qn("w:sz"), "6"); e.set(qn("w:color"), "808080")
        borders.append(e)
    tcPr.append(borders)


def _resolve_logo(path, base_dir):
    if not path:
        return None
    p = path if os.path.isabs(path) else os.path.join(base_dir, path)
    if not os.path.exists(p):
        print(f"[CẢNH BÁO] Không thấy logo: {path} -> để ô placeholder.")
        return None
    if p.lower().endswith(".webp"):
        try:
            from PIL import Image
            png = os.path.splitext(p)[0] + "_cover.png"
            Image.open(p).convert("RGBA").save(png)
            return png
        except Exception as e:
            print(f"[CẢNH BÁO] Không convert được logo .webp ({e}); cần Pillow. Bỏ qua {path}.")
            return None
    return p


def parse_cover(text, base_dir):
    """Tách khối <!-- COVER ... --> khỏi nội dung; trả (dict|None, text_còn_lại)."""
    m = COVER_RE.search(text)
    if not m:
        return None, text
    cover = {"_base": base_dir}
    for ln in m.group(1).splitlines():
        ln = ln.split("#", 1)[0].strip()
        if not ln or ":" not in ln:
            continue
        k, v = ln.split(":", 1)
        cover[k.strip().lower()] = v.strip()
    return cover, text[:m.start()] + text[m.end():]


def build_cover(doc, cover):
    """Chèn trang bìa vào ĐẦU doc (giữ style/lề/footer body). Ngày ghim chân trang bìa."""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    body = doc.element.body
    base = cover.get("_base", ".")

    drop = str(cover.get("drop-title", "")).lower() in ("1", "true", "yes", "có", "co")
    if drop:  # bỏ khối tiêu đề/metadata đầu body tới mục đánh số đầu tiên
        while len(body) > 1:
            first = body[0]
            if first.tag == qn("w:tbl"):
                break
            if first.tag == qn("w:p"):
                txt = "".join(t.text or "" for t in first.iter(qn("w:t")))
                if re.match(r"^\s*\d+\.", txt):
                    break
            body.remove(first)

    anchor = body[0]
    def before(el):
        anchor.addprevious(el)
    def spacer(n=1):
        for _ in range(n):
            p = doc.add_paragraph(); before(p._p)

    # hàng 2 logo (bảng không viền, căn giữa)
    spacer(2)
    left = _resolve_logo(cover.get("logo-left"), base)
    right = _resolve_logo(cover.get("logo-right"), base)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (logo, label) in enumerate([(left, "Logo trái"), (right, "Logo phải")]):
        cell = tbl.rows[0].cells[idx]
        cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if logo:
            cp.add_run().add_picture(logo, width=Cm(4.5))
        else:
            _set_cell_dashed_border(cell)
            r = cp.add_run(f"[ {label} ]"); r.italic = True
            r.font.size = Pt(11); r.font.color.rgb = COVER_GREY
    before(tbl._tbl)

    # tên công ty
    if cover.get("company"):
        spacer(3)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(cover["company"]); r.bold = True
        r.font.size = Pt(22); r.font.color.rgb = COVER_NAVY
        before(p._p)

    # tên tài liệu (mỗi phần | một dòng)
    spacer(6)
    parts = [s.strip() for s in cover.get("title", "").split("|") if s.strip()]
    for i, line in enumerate(parts):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line); r.bold = True; r.font.color.rgb = COVER_NAVY
        r.font.size = Pt(19)
        before(p._p)

    # ngắt sang trang nội dung
    p = doc.add_paragraph(); br = p.add_run()
    brk = OxmlElement("w:br"); brk.set(qn("w:type"), "page"); br._r.append(brk)
    before(p._p)

    # ngày: ghim ở chân TRANG BÌA (footer riêng trang đầu) -> luôn ở đáy trang 1
    if cover.get("date"):
        sec = doc.sections[0]
        sec.different_first_page_header_footer = True
        ff = sec.first_page_footer
        ff.is_linked_to_previous = False
        fp = ff.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in list(fp.runs):
            r._r.getparent().remove(r._r)
        r = fp.add_run(f"Bổ sung cập nhật ngày {cover['date']}")
        r.bold = True; r.font.size = Pt(11); r.font.color.rgb = COVER_NAVY


# ---------- main convert ----------
def convert(md_path, out_path):
    md_dir = os.path.dirname(os.path.abspath(md_path))
    with open(md_path, encoding="utf-8") as f:
        text = f.read()
    cover, text = parse_cover(text, md_dir)
    lines = text.split("\n")

    doc = Document()
    setup_document(doc)
    mmdc = shutil.which("mmdc")

    i = 0
    skip_section_level = None  # đang bỏ qua section tracking
    pending_banner = None      # tiêu đề chờ gắn làm banner cho bảng ngay dưới
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        # heading?
        hm = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if hm:
            level = len(hm.group(1))
            text = hm.group(2).strip()
            # bắt đầu vùng bỏ qua: CHỈ heading "MỤC LỤC KẾ HOẠCH" (bảng tracking).
            # KHÔNG match theo chữ "TRACKING" chung — tiêu đề/section có thể chứa
            # "TRACKING/THEO DÕI TIẾN ĐỘ" sẽ bị cắt nhầm cả phần sau.
            if "MỤC LỤC KẾ HOẠCH" in text.upper():
                skip_section_level = level
                i += 1
                continue
            if skip_section_level is not None and level <= skip_section_level:
                skip_section_level = None  # kết thúc vùng bỏ qua
            if skip_section_level is not None:
                i += 1
                continue
            # Bảng boilerplate (trang ký, quản lý thay đổi): heading đứng NGAY trên bảng
            # -> dùng làm DÒNG BANNER tiêu đề gộp của bảng. KHÔNG áp cho bảng nội dung.
            _bt = text.upper()
            if _bt.startswith("TRANG KÝ") or _bt.startswith("QUẢN LÝ THAY ĐỔI"):
                j = i + 1
                while j < len(lines) and lines[j].strip() == "":
                    j += 1
                if j < len(lines) and lines[j].strip().startswith("|"):
                    pending_banner = text
                    i += 1
                    continue
            add_heading(doc, text, level)
            i += 1
            continue

        if skip_section_level is not None:
            i += 1
            continue

        # marker ngắt trang <!-- PAGEBREAK --> / bỏ qua comment HTML khác
        if stripped.startswith("<!--"):
            if "PAGEBREAK" in stripped.upper():
                p = doc.add_paragraph(); r = p.add_run()
                br = OxmlElement("w:br"); br.set(qn("w:type"), "page"); r._r.append(br)
            i += 1
            continue

        # mermaid fence
        if stripped.startswith("```mermaid"):
            code = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i]); i += 1
            i += 1  # bỏ ```
            render_mermaid("\n".join(code), doc, mmdc, md_dir=md_dir)
            continue

        # fence khác -> giữ như đoạn code đơn giản
        if stripped.startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                p = doc.add_paragraph(); _run(p, lines[i], False, False, None); i += 1
            i += 1
            continue

        # table block
        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not is_table_sep(lines[i]):
                    block.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, block, banner=pending_banner)
            pending_banner = None
            continue

        # blank
        if stripped == "":
            i += 1
            continue

        # ảnh markdown
        im = IMG_RE.search(stripped)
        if im and os.path.exists(im.group(1)):
            try:
                doc.add_picture(im.group(1))
            except Exception:
                doc.add_paragraph(stripped)
            i += 1
            continue

        # placeholder màn hình
        if MH_RE.search(stripped):
            p = doc.add_paragraph()
            _run(p, "[Ảnh màn hình: " + stripped + " — chèn screenshot thật theo ID]",
                 False, True, RGBColor(0x80, 0x80, 0x80))
            i += 1
            continue

        # bullet
        bm = re.match(r"^[-*]\s+(.*)$", stripped)
        if bm:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, strip_traceback(bm.group(1)))
            i += 1
            continue
        nm = re.match(r"^\d+[.)]\s+(.*)$", stripped)
        if nm:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, strip_traceback(nm.group(1)))
            i += 1
            continue

        # blockquote — gom block liên tiếp, thụt lề 1cm, không italic, không space thừa
        if stripped.startswith(">"):
            while i < len(lines) and lines[i].strip().startswith(">"):
                bq_line = lines[i].strip().lstrip(">").strip()
                if bq_line == "":
                    i += 1
                    continue  # bỏ dòng trống trong blockquote, không tạo khoảng trắng
                elif re.match(r"^[-*]\s+(.*)$", bq_line):
                    bm2 = re.match(r"^[-*]\s+(.*)$", bq_line)
                    p = doc.add_paragraph(style="List Bullet")
                    p.paragraph_format.left_indent = Cm(1.5)
                    p.paragraph_format.first_line_indent = Cm(-0.5)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(2)
                    add_inline(p, strip_traceback(bm2.group(1)))
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(1)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(2)
                    add_inline(p, strip_traceback(bq_line))
                i += 1
            continue

        # đoạn văn thường
        p = doc.add_paragraph()
        clean = strip_traceback(stripped)
        add_inline(p, clean)
        # Tiêu đề bước (toàn dòng in đậm, bắt đầu bằng **Bước) → space trước lớn, sau nhỏ
        if re.match(r"^\*\*Bước\s+\d+", stripped):
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(2)
        i += 1

    if cover:
        build_cover(doc, cover)
    doc.save(out_path)
    print(f"[OK] Đã tạo: {out_path}" + ("  (kèm trang bìa)" if cover else ""))
    if not mmdc:
        print("[i] Không thấy mermaid-cli (mmdc) — lưu đồ để placeholder. "
              "Cài: npm i -g @mermaid-js/mermaid-cli, hoặc chèn ảnh từ drawio.")


def main():
    args = [a for a in sys.argv[1:]]
    if not args:
        print('Dùng: python md_to_docx.py "_TOBE-BLUEPRINT.md" [-o out.docx]')
        sys.exit(1)
    md_path = args[0]
    out_path = None
    if "-o" in args:
        out_path = args[args.index("-o") + 1]
    if not os.path.isfile(md_path):
        print(f"[LỖI] Không thấy file: {md_path}")
        sys.exit(1)
    if not out_path:
        out_path = os.path.splitext(md_path)[0] + ".docx"
    convert(md_path, out_path)


if __name__ == "__main__":
    main()
